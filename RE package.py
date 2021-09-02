import re
import pandas as pd
import os
from docx import Document

#查看文件内表格数（一般为3）print(len(document.tables))
#查看表格行列  问题很大 print(len(basic_info.rows),len(basic_info.columns))
def info(docx_file):
    document = Document(docx_file)
    basic_info = document.tables[0]
    print(len(basic_info.rows), len(basic_info.columns))
#先写了姓名和ID住院号，但是我们原始数据不是表格，findall函数也查找不到，应该如何处理？
    return dict(

    住院号 = basic_info.cell(0,0).text,
    姓名 = basic_info.cell(0,2).text,
  #  ID号 = basic_info.cell(0,1).text
    )
print(info('C:\Users\hzf\Desktop.0006893719_朱芳平_2018-04-07.docx'))
#下面是储存xlxs文件
columns= None
datas = []

for file in os.listdir('D:/pycharm/pythonProject'):           #自动寻找文件路径中以.docx结尾的文件
    if file.endswith('.docx'):
        file_path = f'D:/pycharm/pythonProject/{file}'
        #print('病例表',file_path)
        data = info(file_path)
        if not columns:
            columns = list(data.keys())

        datas.append([data[column] for column in columns])

    df = pd.DataFrame(datas,columns = columns)
    df.to_excel('病例表.xlsx',index=False)